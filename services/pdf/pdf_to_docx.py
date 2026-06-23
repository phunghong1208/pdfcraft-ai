"""PDF → DOCX — stack permissive (MIT/BSD): pdfplumber + python-docx + pypdf + pikepdf."""

from __future__ import annotations

import logging
import re
import shutil
import tempfile
import unicodedata
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Literal

import pikepdf
import pdfplumber
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from pdf2image import convert_from_path
from pypdf import PdfReader, PdfWriter

logger = logging.getLogger("pdfcraft.pdf.pdf_to_docx")

ENGINE_EDITABLE = "pdfplumber-docx"
ENGINE_FIXED = "pdfplumber-fixed-layout"
ENGINE_PRESERVE = "preserve-layout-png"
DocxMode = Literal["auto", "preserve_layout", "fixed_layout", "editable"]
PRESERVE_LAYOUT_DPI_DEFAULT = 220
PRESERVE_LAYOUT_DPI_MIN = 150
PRESERVE_LAYOUT_DPI_MAX = 250

# Tiêu đề mục La Mã — không nên nằm trong ô bảng
_ROMAN_HEADING_RE = re.compile(r"^\s*[IVXLCDM]{1,8}[.)]\s+\S")

# Giữ tương thích import cũ
ENGINE_NAME = ENGINE_EDITABLE


@dataclass
class StyledRun:
    text: str
    bold: bool = False
    italic: bool = False
    size: float = 11.0
    fontname: str = ""
    color_rgb: tuple[int, int, int] | None = None


def _color_to_rgb(color: Any) -> tuple[int, int, int] | None:
    """Convert pdfplumber non_stroking_color → (R,G,B) 0-255, None if black/default."""
    if color is None:
        return None
    if not isinstance(color, (list, tuple)) or not color:
        return None
    try:
        if len(color) == 1:
            v = int(float(color[0]) * 255)
            return None if v < 15 else (v, v, v)
        if len(color) == 3:
            r, g, b = (int(float(c) * 255) for c in color)
            return None if (r < 15 and g < 15 and b < 15) else (r, g, b)
        if len(color) == 4:
            c, m, y, k = (float(x) for x in color)
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return None if (r < 15 and g < 15 and b < 15) else (r, g, b)
    except (TypeError, ValueError):
        pass
    return None


def _is_white_color(color: Any) -> bool:
    if color is None:
        return False
    if not isinstance(color, (list, tuple)) or not color:
        return False
    try:
        return all(float(c) > 0.95 for c in color)
    except (TypeError, ValueError):
        return False


def _find_visible_tables(page: Any) -> list[Any]:
    """find_tables() theo đường kẻ thật + lọc bảng giả (line ngang, text bị tách cột)."""
    settings = {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "snap_tolerance": 3,
        "join_tolerance": 3,
    }
    try:
        raw = page.find_tables(table_settings=settings)
    except Exception:
        raw = page.find_tables()
    if isinstance(raw, list):
        all_tables = raw
    else:
        all_tables = list(getattr(raw, "tables", None) or page.find_tables() or [])

    page_w = float(getattr(page, "width", None) or 612)
    page_rects = page.rects or []
    page_lines = page.lines or []
    result: list[Any] = []
    for t in all_tables:
        if not _is_real_table(t, page_w, page):
            continue
        bx0, by0, bx1, by1 = t.bbox
        has_visible = False
        for r in page_rects:
            stroke = r.get("stroking_color")
            fill = r.get("non_stroking_color")
            if _is_white_color(stroke) and _is_white_color(fill):
                continue
            rx0 = float(r.get("x0", 0))
            ry0 = float(r.get("top", 0))
            rx1 = float(r.get("x1", 0))
            ry1 = float(r.get("bottom", 0))
            if rx0 < bx1 + 2 and rx1 > bx0 - 2 and ry0 < by1 + 2 and ry1 > by0 - 2:
                has_visible = True
                break
        if not has_visible:
            for line in page_lines:
                lx0 = float(line.get("x0", 0))
                ly0 = float(line.get("top", 0))
                lx1 = float(line.get("x1", 0))
                ly1 = float(line.get("bottom", 0))
                if lx0 < bx1 + 2 and lx1 > bx0 - 2 and ly0 < by1 + 2 and ly1 > by0 - 2:
                    has_visible = True
                    break
        if has_visible:
            result.append(t)
    # Fallback: không có rect/line rõ nhưng table candidate vẫn hợp lý.
    # Tránh mất bảng thật ở PDF vẽ grid kiểu đặc thù.
    if not result:
        result = [t for t in all_tables if _is_real_table(t, page_w, page)]
    return result


def _table_fill_stats(table: Any) -> tuple[int, int, int, int, float, float]:
    data = table.extract() or []
    rows = len(data)
    cols = max((len(r) for r in data), default=0)
    total = sum(len(r) for r in data)
    filled = sum(1 for r in data for c in r if c and str(c).strip())
    bbox = table.bbox
    w = float(bbox[2]) - float(bbox[0])
    h = float(bbox[3]) - float(bbox[1])
    ratio = filled / total if total else 0.0
    return rows, cols, total, filled, w, h, ratio


def _table_has_cell_rects(page: Any, bbox: tuple[float, ...], min_count: int = 4) -> bool:
    """PDF vẽ bảng bằng ô rect (không phải line) — đếm rect trong vùng."""
    x0, y0, x1, y1 = bbox
    count = 0
    for rect in page.rects or []:
        rx0 = float(rect.get("x0", 0))
        ry0 = float(rect.get("top", 0))
        rx1 = float(rect.get("x1", 0))
        ry1 = float(rect.get("bottom", 0))
        rw, rh = rx1 - rx0, ry1 - ry0
        if rw < 10 or rh < 8:
            continue
        cx, cy = (rx0 + rx1) / 2, (ry0 + ry1) / 2
        if x0 <= cx <= x1 and y0 <= cy <= y1:
            count += 1
    return count >= min_count


def _table_has_grid_lines(page: Any, bbox: tuple[float, ...]) -> bool:
    """Bảng thật cần có cả đường ngang lẫn dọc trong vùng bbox."""
    x0, y0, x1, y1 = bbox
    span_w = max(1.0, x1 - x0)
    span_h = max(1.0, y1 - y0)
    h_count = 0
    v_count = 0

    def _scan_line(lx0: float, lt: float, lx1: float, lb: float) -> None:
        nonlocal h_count, v_count
        if lt < y0 - 3 or lb > y1 + 3:
            return
        lw = abs(lx1 - lx0)
        lh = abs(lb - lt)
        if lh <= 3 and lw >= span_w * 0.2:
            h_count += 1
        elif lw <= 3 and lh >= span_h * 0.12:
            v_count += 1

    for line in page.lines or []:
        _scan_line(
            float(line.get("x0", 0)),
            float(line.get("top", 0)),
            float(line.get("x1", 0)),
            float(line.get("bottom", 0)),
        )
    for rect in page.rects or []:
        rx0 = float(rect.get("x0", 0))
        ry0 = float(rect.get("top", 0))
        rx1 = float(rect.get("x1", 0))
        ry1 = float(rect.get("bottom", 0))
        rw, rh = rx1 - rx0, ry1 - ry0
        if rw < 1 or rh < 1:
            continue
        if rh <= 4 and rw >= span_w * 0.2:
            h_count += 1
        elif rw <= 4 and rh >= span_h * 0.12:
            v_count += 1

    return h_count >= 2 and v_count >= 1


def _table_has_visible_structure(page: Any, bbox: tuple[float, ...]) -> bool:
    return _table_has_grid_lines(page, bbox) or _table_has_cell_rects(page, bbox)


def _is_real_table(table: Any, page_w: float, page: Any | None = None) -> bool:
    """Loại bảng giả: 1 dòng mỏng, text prose bị tách thành nhiều cột."""
    rows, cols, _total, filled, w, h, ratio = _table_fill_stats(table)
    if rows < 1 or cols < 2 or filled == 0:
        return False
    # Dải ngang mỏng (chỉ là đường kẻ / rule)
    if rows == 1 and h < 24:
        row0 = [str(c).strip().lower() for c in (table.extract() or [[]])[0] if c and str(c).strip()]
        header_hits = sum(1 for t in row0 if t in {"phần", "câu", "nội dung", "điểm"})
        if header_hits >= 3:
            return True
        return False
    # Text 1 dòng bị pdfplumber cắt thành nhiều cột nhỏ (vd. "Câu 2 (4.0 điểm)")
    if rows == 1 and cols >= 4:
        texts = [str(c).strip() for row in (table.extract() or []) for c in row if c and str(c).strip()]
        if texts and all(len(t) < 14 for t in texts) and filled <= cols:
            return False
    # Header đề thi 2 cột (không cần grid đầy đủ)
    if rows == 1 and cols == 2 and ratio >= 0.45 and w > page_w * 0.55:
        return True
    # Header đáp án nhiều dòng (ĐỀ CHÍNH THỨC | GỢI Ý ĐÁP ÁN)
    if rows <= 3 and h < 110 and cols >= 2:
        row0 = (table.extract() or [[]])[0]
        row0_text = [str(c).strip() for c in row0 if c and str(c).strip()]
        if len(row0_text) >= 2:
            return True
    # Bảng danh sách nhiều dòng (hospital list, phụ lục, v.v.)
    if rows >= 8 and cols >= 3 and w > page_w * 0.55:
        return True
    # Bảng chấm điểm / rubric
    if rows >= 3 and cols >= 2 and ratio >= 0.08:
        if page is not None and not _table_has_visible_structure(page, table.bbox):
            return False
        return True
    if rows >= 2 and ratio >= 0.18 and cols >= 3:
        if page is not None and not _table_has_visible_structure(page, table.bbox):
            return False
        return True
    return False


def _split_table_rows(
    table: Any,
) -> tuple[list[list[str]], list[str], tuple[float, float, float, float]]:
    """Giữ hàng grid thật; đẩy tiêu đề mục (I. LƯU Ý...) ra ngoài bảng."""
    data = table.extract() or []
    cleaned = [
        [unicodedata.normalize("NFC", str(c or "")).strip() for c in row]
        for row in data if row
    ]
    kept: list[list[str]] = []
    overflow: list[str] = []
    kept_row_indices: list[int] = []

    for ri, row in enumerate(cleaned):
        parts = [c for c in row if c]
        joined = re.sub(r"\s+", " ", " ".join(parts)).strip()
        if not joined:
            continue
        compact = joined.replace(" . ", ". ").replace(" .", ".")
        if _ROMAN_HEADING_RE.match(compact) and len(compact) < 80:
            overflow.append(compact)
            continue
        kept.append(row)
        kept_row_indices.append(ri)

    bbox = tuple(float(x) for x in table.bbox)
    if kept_row_indices and len(kept_row_indices) < len(cleaned):
        try:
            tops: list[float] = []
            bots: list[float] = []
            for ri in kept_row_indices:
                if ri < len(table.rows):
                    for cell in table.rows[ri].cells or []:
                        if cell:
                            tops.append(float(cell[1]))
                            bots.append(float(cell[3]))
            if tops and bots:
                bbox = (bbox[0], min(tops), bbox[2], max(bots))
        except Exception:
            pass

    return kept, overflow, bbox


@dataclass
class TextBlock:
    runs: list[StyledRun]
    top: float
    x0: float = 0.0
    x1: float = 0.0
    bottom: float = 0.0
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT

    @property
    def text(self) -> str:
        return "".join(r.text for r in self.runs)


@dataclass
class TableBlock:
    rows: list[list[str]]
    top: float


@dataclass
class ImageBlock:
    img_bytes: bytes
    ext: str
    top: float
    x0: float
    x1: float
    bottom: float


@dataclass
class HLineBlock:
    top: float
    x0: float
    x1: float
    linewidth: float


@dataclass
class PageContent:
    width: float
    height: float = 792.0
    blocks: list[TextBlock | TableBlock | ImageBlock | HLineBlock] = field(default_factory=list)


def _font_flags(font_name: str) -> tuple[bool, bool]:
    f = (font_name or "").lower().replace("-", "").replace("_", "")
    bold = any(k in f for k in ("bold", "heavy", "black", "demi", "semibold", "extrabold"))
    italic = any(k in f for k in ("italic", "oblique", "ital"))
    if not bold and re.search(r"(^|[^a-z])bd([^a-z]|$)", f):
        bold = True
    if not italic and re.search(r"(^|[^a-z])it([^a-z]|$)", f):
        italic = True
    if not bold:
        m = re.search(r"(\d{3})wght", f)
        if m and int(m.group(1)) >= 600:
            bold = True
    return bold, italic


_FONT_MAP: dict[str, str] = {
    "arial": "Arial",
    "helvetica": "Arial",
    "timesnewroman": "Times New Roman",
    "times": "Times New Roman",
    "courier": "Courier New",
    "couriernew": "Courier New",
    "georgia": "Georgia",
    "verdana": "Verdana",
    "calibri": "Calibri",
    "cambria": "Cambria",
    "tahoma": "Tahoma",
    "trebuchet": "Trebuchet MS",
    "garamond": "Garamond",
    "palatino": "Palatino Linotype",
    "bookantiqua": "Book Antiqua",
}


def _map_font_name(pdf_fontname: str) -> str | None:
    clean = re.sub(r"[-_,.]", "", (pdf_fontname or "").lower())
    clean = re.sub(
        r"(bold|italic|oblique|regular|mt|psmt|ps|it|bd|lt|bk|med|semi|extra|heavy|black|light|thin|condensed|narrow|roman|cyr)",
        "",
        clean,
    ).strip()
    for key, val in _FONT_MAP.items():
        if key in clean:
            return val
    return None


def _align_from_bbox(x0: float, x1: float, page_w: float) -> WD_ALIGN_PARAGRAPH:
    """Prefer LEFT by default; only center short heading-like lines."""
    cx = (x0 + x1) / 2
    width = max(1.0, x1 - x0)
    # Chỉ center khi dòng tương đối ngắn và nằm gần giữa trang.
    if width < page_w * 0.45 and page_w * 0.30 <= cx <= page_w * 0.70:
        if x0 > page_w * 0.10 and x1 < page_w * 0.90:
            return WD_ALIGN_PARAGRAPH.CENTER
    if cx >= page_w * 0.80 and width < page_w * 0.45:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _char_is_rotated(char: dict[str, Any]) -> bool:
    if char.get("upright") is False:
        return True
    matrix = char.get("matrix")
    if matrix and len(matrix) >= 4:
        if abs(float(matrix[1])) > 0.05 or abs(float(matrix[2])) > 0.05:
            return True
    w = float(char.get("width") or 0)
    h = float(char.get("height") or 0)
    text = (char.get("text") or "").strip()
    if text and w > 5 and h > w * 3.5:
        return True
    return False


def _find_rotated_columns(page: Any) -> list[float]:
    """Scan page.chars for rotated characters in margin areas only."""
    page_w = float(page.width or 612)
    margin = 55
    x_vals: list[float] = []
    for char in page.chars or []:
        x0 = float(char.get("x0", 0))
        if x0 > margin and x0 < page_w - margin:
            continue
        if char.get("upright") is False:
            x_vals.append(x0)
            continue
        matrix = char.get("matrix")
        if matrix and len(matrix) >= 4:
            if abs(float(matrix[1])) > 0.05 or abs(float(matrix[2])) > 0.05:
                x_vals.append(x0)
    if not x_vals:
        return []
    x_vals.sort()
    columns = [x_vals[0]]
    for x in x_vals[1:]:
        if x - columns[-1] > 15:
            columns.append(x)
    return columns


def _is_rotated_fragment(word: dict[str, Any], rotated_cols: list[float]) -> bool:
    """Check if word is part of rotated margin text."""
    x0 = float(word["x0"])
    x1 = float(word["x1"])
    w = float(word["x1"]) - x0
    text = (word.get("text") or "").strip()
    if rotated_cols and w < 18 and len(text) <= 3:
        if any(abs(x0 - col) < 8 for col in rotated_cols):
            return True
    # Không loại ký tự đơn trong thân bài để tránh rụng dấu tiếng Việt.
    # Chỉ dựa vào cột xoay thật sự ở lề trang.
    return False


def _point_in_rect(px: float, py: float, rect: tuple[float, ...]) -> bool:
    x0, top, x1, bottom = rect
    return x0 <= px <= x1 and top <= py <= bottom


def _rects_overlap(a: tuple[float, ...], b: tuple[float, ...]) -> bool:
    return a[0] < b[2] and b[0] < a[2] and a[1] < b[3] and b[1] < a[3]


def _group_words_to_lines(words: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    if not words:
        return []
    sorted_words = sorted(words, key=lambda w: (float(w["top"]), float(w["x0"])))
    groups: list[list[dict[str, Any]]] = []
    for word in sorted_words:
        size = float(word.get("size") or 11)
        cy = (float(word["top"]) + float(word["bottom"])) / 2
        tol = max(4.0, size * 0.55)
        placed = False
        for group in groups:
            ref = group[0]
            ref_cy = (float(ref["top"]) + float(ref["bottom"])) / 2
            if abs(cy - ref_cy) <= tol:
                group.append(word)
                placed = True
                break
        if not placed:
            groups.append([word])
    for group in groups:
        group.sort(key=lambda w: float(w["x0"]))
    return groups


def _line_to_text_block(group: list[dict[str, Any]], page_w: float) -> TextBlock | None:
    raw_runs: list[StyledRun] = []
    for i, word in enumerate(group):
        text = unicodedata.normalize("NFC", word.get("text") or "")
        if not text:
            continue
        fontname = str(word.get("fontname") or "")
        size = float(word.get("size") or 11)
        bold, italic = _font_flags(fontname)
        rgb = _color_to_rgb(word.get("non_stroking_color"))
        if i > 0 and raw_runs:
            prev = group[i - 1]
            gap = float(word["x0"]) - float(prev["x1"])
            if gap > max(size * 0.05, 0.5) and raw_runs[-1].text and not raw_runs[-1].text.endswith(" "):
                raw_runs.append(StyledRun(" ", size=size, bold=bold, italic=italic, fontname=fontname, color_rgb=rgb))
        raw_runs.append(StyledRun(text, bold=bold, italic=italic, size=size, fontname=fontname, color_rgb=rgb))
    if not raw_runs:
        return None
    runs: list[StyledRun] = [raw_runs[0]]
    for r in raw_runs[1:]:
        prev = runs[-1]
        if prev.bold == r.bold and prev.italic == r.italic and prev.color_rgb == r.color_rgb and abs(prev.size - r.size) < 0.5:
            prev.text += r.text
        else:
            runs.append(r)
    if not runs:
        return None
    x0 = min(float(w["x0"]) for w in group)
    x1 = max(float(w["x1"]) for w in group)
    top = min(float(w["top"]) for w in group)
    bottom = max(float(w["bottom"]) for w in group)
    return TextBlock(
        runs=runs,
        top=top,
        x0=x0,
        x1=x1,
        bottom=bottom,
        align=_align_from_bbox(x0, x1, page_w),
    )


def _pypdf_page_fallback(pdf_path: Path, page_index: int) -> list[TextBlock]:
    try:
        reader = PdfReader(str(pdf_path))
        if page_index >= len(reader.pages):
            return []
        raw = reader.pages[page_index].extract_text() or ""
    except Exception as exc:
        logger.warning("pypdf fallback failed page %d: %s", page_index + 1, exc)
        return []
    blocks: list[TextBlock] = []
    y = float(page_index) * 1000.0
    for para in re.split(r"\n{2,}", raw.strip()):
        for ln in [l.strip() for l in para.split("\n") if l.strip()]:
            blocks.append(TextBlock(runs=[StyledRun(ln)], top=y, align=WD_ALIGN_PARAGRAPH.LEFT))
            y += 14.0
    return blocks


def _build_image_map(pdf_path: Path) -> dict[int, dict[str, tuple[bytes, str]]]:
    """Pre-extract all images from PDF using pikepdf. Returns {page_idx: {xobj_name: (bytes, ext)}}."""
    result: dict[int, dict[str, tuple[bytes, str]]] = {}
    try:
        with pikepdf.Pdf.open(pdf_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_images: dict[str, tuple[bytes, str]] = {}
                try:
                    resources = page.get("/Resources")
                    if not resources or "/XObject" not in resources:
                        continue
                    for xobj_name, xobj in resources["/XObject"].items():
                        try:
                            name = str(xobj_name).lstrip("/")
                            if xobj.get("/Subtype") != pikepdf.Name("/Image"):
                                continue
                            pdfimg = pikepdf.PdfImage(xobj)
                            pil_img = pdfimg.as_pil_image()
                            if pil_img.mode in ("CMYK", "P", "LA", "PA"):
                                pil_img = pil_img.convert("RGBA" if pil_img.mode == "PA" else "RGB")
                            buf = BytesIO()
                            fmt = "PNG" if pil_img.mode == "RGBA" else "JPEG"
                            pil_img.save(buf, fmt, quality=90)
                            page_images[name] = (buf.getvalue(), fmt.lower())
                        except Exception as exc:
                            logger.debug("image decode %s p%d: %s", xobj_name, page_idx, exc)
                except Exception as exc:
                    logger.debug("page xobjects p%d: %s", page_idx, exc)
                if page_images:
                    result[page_idx] = page_images
    except Exception as exc:
        logger.warning("pikepdf image map failed: %s", exc)
    return result


def _match_page_images(
    page: pdfplumber.page.Page,
    page_idx: int,
    image_map: dict[int, dict[str, tuple[bytes, str]]],
    table_rects: list[tuple[float, float, float, float]],
) -> list[ImageBlock]:
    """Match pdfplumber image positions with pikepdf-extracted image data."""
    plumber_images = page.images or []
    if not plumber_images:
        return []
    page_data = image_map.get(page_idx, {})
    if not page_data:
        return []

    blocks: list[ImageBlock] = []
    for img_info in plumber_images:
        x0, top = float(img_info["x0"]), float(img_info["top"])
        x1, bottom = float(img_info["x1"]), float(img_info["bottom"])
        w, h = x1 - x0, bottom - top
        if w < 8 or h < 8:
            continue
        cx, cy = (x0 + x1) / 2, (top + bottom) / 2
        if any(_point_in_rect(cx, cy, rect) for rect in table_rects):
            continue

        name = str(img_info.get("name", "")).lstrip("/")
        if name not in page_data:
            for candidate in page_data:
                if candidate.startswith(name) or name.startswith(candidate):
                    name = candidate
                    break
            else:
                continue

        data, ext = page_data[name]
        blocks.append(ImageBlock(img_bytes=data, ext=ext, top=top, x0=x0, x1=x1, bottom=bottom))
    return blocks


def _is_watermark_image(
    img: ImageBlock,
    text_blocks: list[TextBlock],
    page_w: float,
    page_h: float,
    table_rects: list[tuple[float, float, float, float]] | None = None,
) -> bool:
    img_area = (img.x1 - img.x0) * (img.bottom - img.top)
    page_area = page_w * page_h
    if page_area <= 0:
        return False
    ratio = img_area / page_area
    if ratio > 0.15:
        return True
    if ratio > 0.02:
        img_rect = (img.x0, img.top, img.x1, img.bottom)
        overlap = sum(
            1
            for tb in text_blocks
            if _rects_overlap((tb.x0, tb.top, tb.x1, tb.bottom), img_rect)
        )
        if table_rects:
            overlap += sum(
                1 for tr in table_rects if _rects_overlap(tr, img_rect)
            )
        if overlap > 1:
            return True
    return False


def _extract_hlines(
    page: pdfplumber.page.Page,
    table_rects: list[tuple[float, float, float, float]],
) -> list[HLineBlock]:
    """Extract horizontal separator lines from PDF page."""
    page_w = float(page.width or 612)
    min_span = page_w * 0.35
    blocks: list[HLineBlock] = []

    for line in page.lines or []:
        x0, y0 = float(line["x0"]), float(line["top"])
        x1, y1 = float(line["x1"]), float(line["bottom"])
        if abs(y1 - y0) > 3:
            continue
        if (x1 - x0) < min_span:
            continue
        lw = float(line.get("linewidth", 1) or 1)
        cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
        if any(_point_in_rect(cx, cy, rect) for rect in table_rects):
            continue
        blocks.append(HLineBlock(top=y0, x0=x0, x1=x1, linewidth=lw))

    for rect in page.rects or []:
        x0, y0 = float(rect["x0"]), float(rect["top"])
        x1, y1 = float(rect["x1"]), float(rect["bottom"])
        w, h = x1 - x0, y1 - y0
        if h > 4 or w < min_span:
            continue
        cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
        if any(_point_in_rect(cx, cy, tr) for tr in table_rects):
            continue
        blocks.append(HLineBlock(top=y0, x0=x0, x1=x1, linewidth=max(h, 0.5)))

    return blocks


def _is_decorative_hline(hline: HLineBlock, text_blocks: list[TextBlock]) -> bool:
    """Bỏ đường kẻ ngang giữa các dòng thơ / đoạn văn — không phải viền bảng."""
    for tb in text_blocks:
        if hline.x0 < tb.x1 - 4 and hline.x1 > tb.x0 + 4:
            if abs(hline.top - tb.top) < 18 or abs(hline.top - tb.bottom) < 18:
                return True
            if tb.top < hline.top < tb.bottom:
                return True
    return False


def _should_merge_lines(current: TextBlock, next_block: TextBlock, page_w: float = 0) -> bool:
    cur_sizes = [r.size for r in current.runs if r.text.strip()]
    next_sizes = [r.size for r in next_block.runs if r.text.strip()]
    if not cur_sizes or not next_sizes:
        return False
    cur_size = sum(cur_sizes) / len(cur_sizes)
    next_size = sum(next_sizes) / len(next_sizes)
    if abs(cur_size - next_size) > 1.5:
        return False
    gap = next_block.top - current.bottom
    if gap < -2 or gap > cur_size * 1.5:
        return False
    if abs(next_block.x0 - current.x0) > cur_size * 4:
        return False
    cur_all_bold = all(r.bold for r in current.runs if r.text.strip())
    if cur_all_bold and len(current.text.strip()) < 80:
        return False
    next_all_bold = all(r.bold for r in next_block.runs if r.text.strip())
    if next_all_bold and len(next_block.text.strip()) < 120:
        return False
    if page_w > 0 and current.x1 < page_w * 0.72:
        return False
    next_text = next_block.text.strip()
    if re.match(r"^(Câu|câu|Bài|bài)\s+\d", next_text):
        return False
    if re.match(r"^[IVX]+\.\s", next_text):
        return False
    if re.match(r"^\d+\.\s", next_text):
        return False
    return True


def _merge_lines_to_paragraphs(text_blocks: list[TextBlock], page_w: float) -> list[TextBlock]:
    if len(text_blocks) <= 1:
        return text_blocks
    text_blocks.sort(key=lambda b: b.top)
    merged: list[TextBlock] = []
    current = text_blocks[0]
    for i in range(1, len(text_blocks)):
        nb = text_blocks[i]
        if _should_merge_lines(current, nb, page_w):
            cur_size = sum(r.size for r in current.runs if r.text.strip()) / max(1, sum(1 for r in current.runs if r.text.strip()))
            current.runs.append(StyledRun(" ", size=cur_size))
            current.runs.extend(nb.runs)
            current.bottom = nb.bottom
            current.x1 = max(current.x1, nb.x1)
        else:
            merged.append(current)
            current = nb
    merged.append(current)
    for block in merged:
        block.align = _align_from_bbox(block.x0, block.x1, page_w)
    return merged


def _median_font_size(blocks: list[TextBlock]) -> float:
    sizes: list[float] = []
    for b in blocks:
        for r in b.runs:
            if r.text.strip() and r.size > 0:
                sizes.append(r.size)
    if not sizes:
        return 11.0
    sizes.sort()
    mid = len(sizes) // 2
    return sizes[mid]


def _detect_heading_level(block: TextBlock, median_size: float) -> int:
    """0 = not heading, 1-3 = heading level."""
    avg_size = sum(r.size for r in block.runs if r.text.strip()) / max(1, sum(1 for r in block.runs if r.text.strip()))
    is_bold = all(r.bold for r in block.runs if r.text.strip())
    text = block.text.strip()
    if len(text) > 200 or not text:
        return 0
    if avg_size >= median_size * 1.6:
        return 1
    if avg_size >= median_size * 1.3:
        return 2
    if is_bold and avg_size >= median_size * 1.05 and len(text) < 120:
        return 3
    return 0


def _extract_page(page: pdfplumber.page.Page, pdf_path: Path, image_map: dict[int, dict[str, tuple[bytes, str]]] | None = None) -> PageContent:
    page_w = float(page.width or 612)
    page_h = float(page.height or 792)
    content = PageContent(width=page_w, height=page_h)
    page_idx = (page.page_number or 1) - 1

    table_rects: list[tuple[float, float, float, float]] = []
    table_blocks: list[TableBlock] = []
    try:
        for table in _find_visible_tables(page):
            cleaned, overflow, bbox = _split_table_rows(table)
            _ = overflow
            if not cleaned or not any(any(cell for cell in row) for row in cleaned):
                continue
            # Dùng bbox đã trim để text hàng tiêu đề (I./II.) vẫn giữ đúng vị trí ngoài bảng.
            table_rects.append(bbox)
            table_blocks.append(TableBlock(rows=cleaned, top=float(bbox[1])))
    except Exception as exc:
        logger.warning("table extract failed page %d: %s", page_idx + 1, exc)

    rotated_cols = _find_rotated_columns(page)

    words = page.extract_words(
        x_tolerance=1.5,
        y_tolerance=2,
        keep_blank_chars=False,
        extra_attrs=["fontname", "size", "non_stroking_color"],
    ) or []

    filtered: list[dict[str, Any]] = []
    for word in words:
        if not (word.get("text") or "").strip():
            continue
        cx = (float(word["x0"]) + float(word["x1"])) / 2
        cy = (float(word["top"]) + float(word["bottom"])) / 2
        if any(_point_in_rect(cx, cy, rect) for rect in table_rects):
            continue
        if _is_rotated_fragment(word, rotated_cols):
            continue
        filtered.append(word)

    text_blocks: list[TextBlock] = []
    for group in _group_words_to_lines(filtered):
        block = _line_to_text_block(group, page_w)
        if block:
            text_blocks.append(block)

    if not text_blocks and not table_blocks:
        text_blocks = _pypdf_page_fallback(pdf_path, page_idx)

    img_blocks: list[ImageBlock] = []
    if image_map is not None:
        raw_imgs = _match_page_images(page, page_idx, image_map, table_rects)
        img_blocks = [
            img for img in raw_imgs
            if not _is_watermark_image(img, text_blocks, page_w, page_h, table_rects)
        ]

    text_blocks = _merge_lines_to_paragraphs(text_blocks, page_w)

    ordered: list[TextBlock | TableBlock | ImageBlock | HLineBlock] = []
    ordered.extend(text_blocks)
    ordered.extend(table_blocks)
    ordered.extend(img_blocks)
    def _block_priority(block: TextBlock | TableBlock | ImageBlock | HLineBlock) -> int:
        # Khi cùng top, text (đặc biệt I./II.) phải đứng trước table.
        if isinstance(block, TextBlock):
            return 0
        if isinstance(block, TableBlock):
            return 1
        if isinstance(block, ImageBlock):
            return 2
        return 3

    ordered.sort(key=lambda b: (b.top, _block_priority(b)))
    content.blocks = ordered
    return content


def _apply_run_style(run, run_data: StyledRun) -> None:
    """Apply bold/italic/size/font/color from StyledRun to DOCX Run."""
    run.bold = run_data.bold
    run.italic = run_data.italic
    if run_data.size > 0:
        run.font.size = Pt(max(6, round(run_data.size)))
    mapped = _map_font_name(run_data.fontname)
    if mapped:
        run.font.name = mapped
    if run_data.color_rgb:
        run.font.color.rgb = RGBColor(*run_data.color_rgb)


def _paragraph_space_before(gap: float, block: TextBlock, prev_kind: str) -> float:
    """Khoảng cách Word tối thiểu — không copy gap tọa độ PDF (gây phình trang)."""
    if prev_kind == "table":
        return 6.0
    if gap <= 0:
        return 0.0
    text = block.text.strip()
    if _ROMAN_HEADING_RE.match(text):
        return 12.0
    sizes = [r.size for r in block.runs if r.text.strip()]
    fs = sum(sizes) / len(sizes) if sizes else 11.0
    if gap > fs * 1.75:
        return 6.0
    return 0.0


def _add_text_block(
    doc: Document,
    block: TextBlock,
    heading_level: int = 0,
    space_before_pt: float = 0,
) -> None:
    if heading_level and 1 <= heading_level <= 3:
        para = doc.add_heading(level=heading_level)
        para.clear()
    else:
        para = doc.add_paragraph()
    para.alignment = block.align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.05
    if space_before_pt > 0:
        para.paragraph_format.space_before = Pt(min(space_before_pt, 12))
    for run_data in block.runs:
        if not run_data.text:
            continue
        run = para.add_run(run_data.text)
        _apply_run_style(run, run_data)


def _set_table_borders(table, sz: int = 4, color: str = "000000") -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _add_table_block(doc: Document, block: TableBlock) -> None:
    if not block.rows:
        return
    cols = max(len(row) for row in block.rows)
    if cols < 1:
        return
    non_empty_cols = [
        j for j in range(cols)
        if any((row[j] if j < len(row) else "").strip() for row in block.rows)
    ]
    if non_empty_cols and len(non_empty_cols) < cols:
        compact_rows = [
            [(row[j] if j < len(row) else "") for j in non_empty_cols]
            for row in block.rows
        ]
    else:
        compact_rows = block.rows

    cols = max(len(row) for row in compact_rows)
    if cols < 1:
        return

    table = doc.add_table(rows=len(compact_rows), cols=cols)
    table.style = "Table Grid"
    _set_table_borders(table)
    for i, row in enumerate(compact_rows):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.0
                for run in para.runs:
                    run.font.size = Pt(9)


def _add_image_block(doc: Document, block: ImageBlock, page_width_pt: float) -> None:
    try:
        stream = BytesIO(block.img_bytes)
        w_pt = block.x1 - block.x0
        h_pt = block.bottom - block.top
        max_w = max(72.0, page_width_pt - 72)
        if w_pt > max_w:
            scale = max_w / w_pt
            w_pt *= scale
            h_pt *= scale
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run()
        run.add_picture(stream, width=Emu(_pt_to_emu(w_pt)), height=Emu(_pt_to_emu(h_pt)))
    except Exception as exc:
        logger.warning("image insert failed: %s", exc)


def _add_hline_block(doc: Document, block: HLineBlock) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    p_pr = para._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom_el = OxmlElement("w:bottom")
    sz = max(4, min(48, round(block.linewidth * 8)))
    bottom_el.set(qn("w:val"), "single")
    bottom_el.set(qn("w:sz"), str(sz))
    bottom_el.set(qn("w:space"), "0")
    bottom_el.set(qn("w:color"), "000000")
    p_bdr.append(bottom_el)
    p_pr.append(p_bdr)


def _build_docx(pages: list[PageContent], output_path: Path) -> None:
    doc = Document()
    # Margin mặc định Word — text chảy tự nhiên, không ép spacing theo PDF
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    for page_idx, page in enumerate(pages):
        if page_idx > 0:
            doc.add_page_break()
        prev_bottom = 0.0
        prev_kind = ""
        for block in page.blocks:
            if isinstance(block, TextBlock):
                gap = max(0.0, block.top - prev_bottom) if prev_bottom > 0 else 0.0
                spacing = _paragraph_space_before(gap, block, prev_kind)
                _add_text_block(doc, block, heading_level=0, space_before_pt=spacing)
                prev_bottom = block.bottom
                prev_kind = "text"
            elif isinstance(block, TableBlock):
                if prev_kind:
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(6)
                _add_table_block(doc, block)
                prev_bottom = block.top + 20
                prev_kind = "table"
            elif isinstance(block, ImageBlock):
                _add_image_block(doc, block, page.width)
                prev_bottom = block.bottom
                prev_kind = "image"
    if not doc.paragraphs and not doc.tables:
        doc.add_paragraph("")
    doc.save(str(output_path))


def _paragraph_has_content(para) -> bool:
    if (para.text or "").strip():
        return True
    for run in para.runs:
        xml = run._element.xml
        if "w:drawing" in xml or "w:pict" in xml or "w:object" in xml:
            return True
    return False


def _has_page_break(para) -> bool:
    for br in para._element.xpath(".//w:br"):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _remove_blank_pages(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    paras = doc.paragraphs

    remove_els = []
    i = 0
    while i < len(paras):
        if _has_page_break(paras[i]):
            j = i + 1
            blank_run = True
            while j < len(paras):
                if _paragraph_has_content(paras[j]):
                    blank_run = False
                    break
                if _has_page_break(paras[j]):
                    break
                remove_els.append(paras[j]._element)
                j += 1
            if blank_run and j < len(paras) and _has_page_break(paras[j]):
                remove_els.append(paras[j]._element)
        i += 1

    for el in remove_els:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    paras = doc.paragraphs
    prev_had_break = False
    for para in paras:
        br = _has_page_break(para)
        if br and prev_had_break and not _paragraph_has_content(para):
            for br_el in para._element.xpath(".//w:br"):
                if br_el.get(qn("w:type")) == "page":
                    br_el.getparent().remove(br_el)
            br = False
        prev_had_break = br

    while doc.paragraphs and not _paragraph_has_content(doc.paragraphs[0]):
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    while doc.paragraphs and not _paragraph_has_content(doc.paragraphs[-1]):
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    doc.save(str(docx_path))


def _page_has_text_pypdf(page) -> bool:
    try:
        return bool((page.extract_text() or "").strip())
    except Exception:
        return True


def _filter_empty_pdf_pages(src: Path, dst: Path) -> Path:
    reader = PdfReader(str(src))
    if len(reader.pages) <= 1:
        return src
    kept = [p for p in reader.pages if _page_has_text_pypdf(p)]
    if not kept or len(kept) == len(reader.pages):
        return src
    writer = PdfWriter()
    for p in kept:
        writer.add_page(p)
    writer.write(str(dst))
    logger.info("Filtered %d empty PDF pages (pypdf)", len(reader.pages) - len(kept))
    return dst


def _extract_all_pages(pdf_path: Path) -> list[PageContent]:
    image_map = _build_image_map(pdf_path)
    pages: list[PageContent] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            pages.append(_extract_page(page, pdf_path, image_map))
    return pages


# ── Layout analysis & preserve-layout DOCX ───────────────────────


@dataclass
class PdfLayoutInfo:
    image_count: int = 0
    vector_line_count: int = 0
    has_rotated_text: bool = False
    has_complex_tables: bool = False
    column_count: int = 1
    page_count: int = 0


def _pt_to_emu(pt: float) -> int:
    return int(round(pt * 12700))


def _clamp_preserve_dpi(dpi: int) -> int:
    return max(PRESERVE_LAYOUT_DPI_MIN, min(PRESERVE_LAYOUT_DPI_MAX, dpi))


def _estimate_column_count(page: pdfplumber.page.Page) -> int:
    words = page.extract_words() or []
    if len(words) < 12:
        return 1
    page_w = float(page.width or 612)
    left_words = 0
    right_words = 0
    center_words = 0
    valid_words = 0

    for w in words:
        text = (w.get("text") or "").strip()
        if len(text) <= 1:
            continue
        x0 = float(w.get("x0") or 0)
        x1 = float(w.get("x1") or x0)
        ww = x1 - x0
        if ww < 4:
            continue
        valid_words += 1
        # Chạm dải giữa trang => text một cột hoặc tiêu đề toàn trang
        if x0 < page_w * 0.52 and x1 > page_w * 0.48:
            center_words += 1
            continue
        if x1 <= page_w * 0.48:
            left_words += 1
        elif x0 >= page_w * 0.52:
            right_words += 1
        else:
            center_words += 1

    if valid_words < 60:
        return 1

    # Hai cột thật: cả hai bên đều dày chữ và phần giữa ít.
    if (
        left_words >= 80
        and right_words >= 80
        and center_words <= min(left_words, right_words) * 0.22
    ):
        return 2
    return 1


def _analyze_page_layout(page: pdfplumber.page.Page) -> PdfLayoutInfo:
    info = PdfLayoutInfo(page_count=1)
    try:
        info.image_count = len(page.images or [])
    except Exception:
        info.image_count = 0

    try:
        info.vector_line_count = (
            len(page.lines or [])
            + len(page.rects or [])
            + len(page.curves or [])
        )
    except Exception:
        info.vector_line_count = 0

    try:
        chars = page.chars or []
        info.has_rotated_text = any(_char_is_rotated(c) for c in chars)
    except Exception:
        info.has_rotated_text = False

    try:
        tables = _find_visible_tables(page)
        for table in tables:
            data = table.extract() or []
            rows = len(data)
            cols = max((len(row) for row in data), default=0)
            if rows * cols > 12 or len(tables) > 1:
                info.has_complex_tables = True
                break
    except Exception:
        pass

    try:
        info.column_count = _estimate_column_count(page)
    except Exception:
        info.column_count = 1

    return info


def analyze_pdf_layout(pdf_path: Path) -> PdfLayoutInfo:
    """Phân tích độ phức tạp bố cục để chọn chế độ DOCX."""
    merged = PdfLayoutInfo()
    with pdfplumber.open(pdf_path) as pdf:
        merged.page_count = len(pdf.pages)
        for page in pdf.pages:
            page_info = _analyze_page_layout(page)
            merged.image_count += page_info.image_count
            merged.vector_line_count = max(
                merged.vector_line_count,
                page_info.vector_line_count,
            )
            merged.has_rotated_text = merged.has_rotated_text or page_info.has_rotated_text
            merged.has_complex_tables = (
                merged.has_complex_tables or page_info.has_complex_tables
            )
            merged.column_count = max(merged.column_count, page_info.column_count)
    return merged


def _pt_to_twips(pt: float) -> int:
    return int(round(pt * 20))


def _add_framed_text_block(doc: Document, block: TextBlock) -> None:
    """Đặt đoạn text tại tọa độ PDF (text thật, chỉnh sửa được trong Word)."""
    if not block.text.strip():
        return

    raw_w = block.x1 - block.x0
    width_pt = max(12.0, raw_w + max(6.0, raw_w * 0.10))
    height_pt = max(8.0, block.bottom - block.top)

    para = doc.add_paragraph()
    p_pr = para._element.get_or_add_pPr()
    frame = OxmlElement("w:framePr")
    frame.set(qn("w:w"), str(_pt_to_twips(width_pt)))
    frame.set(qn("w:h"), str(_pt_to_twips(height_pt)))
    frame.set(qn("w:vAnchor"), "page")
    frame.set(qn("w:hAnchor"), "page")
    frame.set(qn("w:x"), str(_pt_to_twips(block.x0)))
    frame.set(qn("w:y"), str(_pt_to_twips(block.top)))
    frame.set(qn("w:hRule"), "atLeast")
    frame.set(qn("w:wrap"), "notBeside")
    p_pr.append(frame)

    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1
    if block.align is not None:
        para.alignment = block.align

    for run_data in block.runs:
        if not run_data.text:
            continue
        run = para.add_run(run_data.text)
        _apply_run_style(run, run_data)


def _extract_positioned_page_blocks(
    page: pdfplumber.page.Page,
    pdf_path: Path,
    table_rects: list[tuple[float, float, float, float]] | None = None,
) -> list[TextBlock]:
    """Trích text + bbox từng dòng — dùng cho fixed-layout DOCX. Bỏ qua vùng table."""
    page_w = float(page.width or 612)
    page_idx = (page.page_number or 1) - 1
    blocks: list[TextBlock] = []

    if table_rects is None:
        table_rects = []
        try:
            for table in _find_visible_tables(page):
                table_rects.append(table.bbox)
        except Exception:
            pass

    rotated_cols = _find_rotated_columns(page)

    words = page.extract_words(
        x_tolerance=1.5,
        y_tolerance=2,
        keep_blank_chars=False,
        extra_attrs=["fontname", "size", "non_stroking_color"],
    ) or []

    filtered: list[dict[str, Any]] = []
    for word in words:
        if not (word.get("text") or "").strip():
            continue
        cx = (float(word["x0"]) + float(word["x1"])) / 2
        cy = (float(word["top"]) + float(word["bottom"])) / 2
        if any(_point_in_rect(cx, cy, rect) for rect in table_rects):
            continue
        if _is_rotated_fragment(word, rotated_cols):
            continue
        filtered.append(word)

    for group in _group_words_to_lines(filtered):
        block = _line_to_text_block(group, page_w)
        if block:
            blocks.append(block)

    if not blocks and not table_rects:
        blocks = _pypdf_page_fallback(pdf_path, page_idx)

    blocks.sort(key=lambda b: (b.top, b.x0))
    return blocks


def _add_anchored_image(doc: Document, block: ImageBlock) -> None:
    """Insert image positioned via framePr in fixed-layout mode."""
    try:
        w_pt = max(12.0, block.x1 - block.x0)
        h_pt = max(8.0, block.bottom - block.top)

        para = doc.add_paragraph()
        p_pr = para._element.get_or_add_pPr()
        frame = OxmlElement("w:framePr")
        frame.set(qn("w:w"), str(_pt_to_twips(w_pt)))
        frame.set(qn("w:h"), str(_pt_to_twips(h_pt)))
        frame.set(qn("w:vAnchor"), "page")
        frame.set(qn("w:hAnchor"), "page")
        frame.set(qn("w:x"), str(_pt_to_twips(block.x0)))
        frame.set(qn("w:y"), str(_pt_to_twips(block.top)))
        frame.set(qn("w:hRule"), "exact")
        frame.set(qn("w:wrap"), "notBeside")
        p_pr.append(frame)

        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        run.add_picture(BytesIO(block.img_bytes), width=Emu(_pt_to_emu(w_pt)), height=Emu(_pt_to_emu(h_pt)))
    except Exception as exc:
        logger.warning("anchored image insert failed: %s", exc)


def _cell_formatting(page: Any, cell_bbox: tuple) -> tuple[bool, bool, tuple[int, int, int] | None]:
    """Detect bold/italic/color from chars within a table cell bbox."""
    if page is None or cell_bbox is None:
        return False, False, None
    try:
        x0, top, x1, bottom = cell_bbox
        cropped = page.crop((x0 + 0.5, top + 0.5, x1 - 0.5, bottom - 0.5))
        chars = cropped.chars or []
        if not chars:
            return False, False, None
        text_chars = [c for c in chars if (c.get("text") or "").strip()]
        if not text_chars:
            return False, False, None
        bold_count = sum(1 for c in text_chars if _font_flags(c.get("fontname") or "")[0])
        italic_count = sum(1 for c in text_chars if _font_flags(c.get("fontname") or "")[1])
        total = len(text_chars)
        is_bold = bold_count > total * 0.6
        is_italic = italic_count > total * 0.6
        first_color = _color_to_rgb(text_chars[0].get("non_stroking_color"))
        return is_bold, is_italic, first_color
    except Exception:
        return False, False, None


def _add_positioned_table(
    doc: Document,
    table_obj: Any,
    page_w: float,
    page: Any = None,
) -> None:
    """Render pdfplumber table as positioned DOCX table with borders."""
    data = table_obj.extract() or []
    cleaned = [[unicodedata.normalize("NFC", str(c or "")).strip() for c in row] for row in data if row]
    if not cleaned or not any(any(cell for cell in row) for row in cleaned):
        return

    cols = max(len(row) for row in cleaned)
    if cols < 1:
        return

    non_empty = [j for j in range(cols) if any((row[j] if j < len(row) else "") for row in cleaned)]
    if 0 < len(non_empty) < cols:
        cleaned = [[(row[j] if j < len(row) else "") for j in non_empty] for row in cleaned]
        cols = len(non_empty)

    bbox = table_obj.bbox
    tbl_x0 = float(bbox[0])
    tbl_top = float(bbox[1])
    tbl_w = float(bbox[2]) - tbl_x0

    table = doc.add_table(rows=len(cleaned), cols=cols)
    _set_table_borders(table)

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_p = OxmlElement("w:tblpPr")
    tbl_p.set(qn("w:vertAnchor"), "page")
    tbl_p.set(qn("w:horzAnchor"), "page")
    tbl_p.set(qn("w:tblpX"), str(_pt_to_twips(tbl_x0)))
    tbl_p.set(qn("w:tblpY"), str(_pt_to_twips(tbl_top)))
    tbl_pr.append(tbl_p)

    max_tbl_w = page_w * 0.92
    actual_w = min(tbl_w, max_tbl_w)
    tbl_w_el = OxmlElement("w:tblW")
    tbl_w_el.set(qn("w:w"), str(_pt_to_twips(actual_w)))
    tbl_w_el.set(qn("w:type"), "dxa")
    for old in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old)
    tbl_pr.append(tbl_w_el)

    tbl_cell_mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "28")
        el.set(qn("w:type"), "dxa")
        tbl_cell_mar.append(el)
    tbl_pr.append(tbl_cell_mar)

    col_widths_pt = [actual_w / cols] * cols
    try:
        cells = getattr(table_obj, "cells", None) or []
        if cells:
            x_edges = sorted(set(float(c[0]) for c in cells) | set(float(c[2]) for c in cells))
            if len(x_edges) == cols + 1:
                col_widths_pt = [x_edges[i + 1] - x_edges[i] for i in range(cols)]
            else:
                max_lens = [0] * cols
                for row in cleaned:
                    for j in range(min(len(row), cols)):
                        max_lens[j] = max(max_lens[j], len(row[j]))
                total_len = sum(max_lens) or 1
                min_w = actual_w * 0.06
                col_widths_pt = [max(actual_w * ml / total_len, min_w) for ml in max_lens]
                scale = actual_w / sum(col_widths_pt)
                col_widths_pt = [w * scale for w in col_widths_pt]
    except Exception:
        pass

    # Get per-cell bboxes for formatting extraction
    row_cell_bboxes: list[list[tuple | None]] = []
    try:
        for row_obj in table_obj.rows:
            row_cells = list(row_obj.cells) if row_obj.cells else [None] * cols
            if non_empty and len(row_cells) > cols:
                row_cells = [row_cells[j] if j < len(row_cells) else None for j in non_empty]
            row_cell_bboxes.append(row_cells)
    except Exception:
        row_cell_bboxes = [[None] * cols for _ in cleaned]

    for i, row in enumerate(cleaned):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            w_val = _pt_to_twips(col_widths_pt[j]) if j < len(col_widths_pt) else _pt_to_twips(actual_w / cols)
            tc_w.set(qn("w:w"), str(w_val))
            tc_w.set(qn("w:type"), "dxa")
            for old in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(old)
            tc_pr.append(tc_w)
            # Apply formatting from PDF chars
            cell_bbox = row_cell_bboxes[i][j] if i < len(row_cell_bboxes) and j < len(row_cell_bboxes[i]) else None
            is_bold, is_italic, cell_color = _cell_formatting(page, cell_bbox)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                for run in para.runs:
                    run.font.size = Pt(8)
                    if is_bold:
                        run.bold = True
                    if is_italic:
                        run.italic = True
                    if cell_color:
                        run.font.color.rgb = RGBColor(*cell_color)


def _build_fixed_layout_docx(pdf_path: Path, output_path: Path) -> None:
    """DOCX text đặt theo tọa độ PDF — chỉnh sửa được, không nhúng ảnh full-page."""
    image_map = _build_image_map(pdf_path)
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            if page_idx > 0:
                doc.add_section(WD_SECTION.NEW_PAGE)

            page_w = float(page.width or 612)
            page_h = float(page.height or 792)
            section = doc.sections[-1]
            section.page_width = Emu(_pt_to_emu(page_w))
            section.page_height = Emu(_pt_to_emu(page_h))
            _zero_section_margins(section)

            table_rects: list[tuple[float, float, float, float]] = []
            tables: list[Any] = []
            try:
                cands: list[tuple[Any, tuple[float, float, float, float]]] = []
                for table in _find_visible_tables(page):
                    cleaned, _overflow, bbox = _split_table_rows(table)
                    if not cleaned or not any(any(cell for cell in row) for row in cleaned):
                        continue
                    tbl_h = float(bbox[3]) - float(bbox[1])
                    if len(cleaned) <= 1 and tbl_h < 25:
                        continue
                    total = sum(len(row) for row in cleaned)
                    filled = sum(1 for row in cleaned for c in row if c)
                    if total > 0 and filled / total > 0.15:
                        capped = (
                            float(bbox[0]),
                            float(bbox[1]),
                            float(bbox[2]),
                            min(float(bbox[3]), page_h),
                        )
                        cands.append((table, capped, cleaned))
                for ci, (tbl, rect, cleaned_rows) in enumerate(cands):
                    nested = any(
                        cj != ci
                        and cands[cj][1][0] <= rect[0]
                        and cands[cj][1][1] <= rect[1]
                        and cands[cj][1][2] >= rect[2]
                        and cands[cj][1][3] >= rect[3]
                        for cj in range(len(cands))
                    )
                    if not nested:
                        table_rects.append(rect)
                        tables.append(tbl)
            except Exception:
                pass

            text_blocks = _extract_positioned_page_blocks(page, pdf_path, table_rects)
            for block in text_blocks:
                _add_framed_text_block(doc, block)

            for table in tables:
                _add_positioned_table(doc, table, page_w, page=page)

            raw_imgs = _match_page_images(page, page_idx, image_map, table_rects)
            for img_block in raw_imgs:
                if not _is_watermark_image(img_block, text_blocks, page_w, page_h, table_rects):
                    _add_anchored_image(doc, img_block)

    doc.save(str(output_path))


def choose_docx_mode(info: PdfLayoutInfo) -> DocxMode:
    """Auto mode: multi-column/complex layout ưu tiên fixed_layout."""
    if info.has_rotated_text:
        return "fixed_layout"
    if info.column_count > 1 and info.page_count > 0:
        return "fixed_layout"
    if info.image_count > 0 and info.vector_line_count > 120:
        return "fixed_layout"
    return "editable"


def _resolve_docx_mode(requested: str, pdf_path: Path) -> tuple[DocxMode, PdfLayoutInfo]:
    layout = analyze_pdf_layout(pdf_path)
    req = (requested or "auto").strip().lower()
    if req in ("preserve_layout", "preserve-layout", "preserve", "image", "snapshot"):
        return "preserve_layout", layout
    if req in ("fixed_layout", "fixed-layout", "positioned", "layout"):
        return "fixed_layout", layout
    if req in ("editable", "text", "flow"):
        return "editable", layout
    return choose_docx_mode(layout), layout


def _pdf_page_sizes_pt(pdf_path: Path) -> list[tuple[float, float]]:
    reader = PdfReader(str(pdf_path))
    sizes: list[tuple[float, float]] = []
    for page in reader.pages:
        mb = page.mediabox
        sizes.append((float(mb.width), float(mb.height)))
    return sizes


def _zero_section_margins(section) -> None:
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)


def _build_preserve_layout_docx(
    pdf_path: Path,
    output_path: Path,
    dpi: int = PRESERVE_LAYOUT_DPI_DEFAULT,
) -> None:
    """Render từng trang PDF → PNG, nhúng full-page vào DOCX (margin 0)."""
    dpi = _clamp_preserve_dpi(dpi)
    page_sizes = _pdf_page_sizes_pt(pdf_path)
    if not page_sizes:
        raise ValueError("PDF has no pages")

    doc = Document()
    scratch = Path(tempfile.mkdtemp(prefix="pdfcraft-docx-raster-"))
    try:
        for page_idx, (page_w, page_h) in enumerate(page_sizes):
            if page_idx > 0:
                doc.add_section(WD_SECTION.NEW_PAGE)

            section = doc.sections[-1]
            section.page_width = Emu(_pt_to_emu(page_w))
            section.page_height = Emu(_pt_to_emu(page_h))
            _zero_section_margins(section)

            page_no = page_idx + 1
            paths = convert_from_path(
                str(pdf_path),
                dpi=dpi,
                first_page=page_no,
                last_page=page_no,
                fmt="png",
                output_folder=str(scratch),
                paths_only=True,
                thread_count=1,
            )
            if not paths:
                raise RuntimeError(f"pdf2image failed for page {page_no}")

            img_path = Path(paths[0])
            try:
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1
                run = para.add_run()
                run.add_picture(
                    str(img_path),
                    width=Emu(_pt_to_emu(page_w)),
                    height=Emu(_pt_to_emu(page_h)),
                )
            finally:
                img_path.unlink(missing_ok=True)
    finally:
        shutil.rmtree(scratch, ignore_errors=True)

    doc.save(str(output_path))


def _build_editable_docx(pdf_path: Path, output_path: Path) -> None:
    work_pdf = pdf_path.parent / "filtered.pdf"
    try:
        src = _filter_empty_pdf_pages(pdf_path, work_pdf)
        pages = _extract_all_pages(src)
        _build_docx(pages, output_path)
        _remove_blank_pages(output_path)
    finally:
        if work_pdf.exists() and work_pdf != pdf_path:
            work_pdf.unlink(missing_ok=True)


def convert_pdf_to_docx(
    input_path: Path,
    output_path: Path,
    mode: str = "auto",
    dpi: int = PRESERVE_LAYOUT_DPI_DEFAULT,
) -> tuple[str, str, PdfLayoutInfo]:
    """
    Chuyển PDF → DOCX.

    mode: auto | preserve_layout | editable
    Trả về (engine, mode_used, layout_info).
    """
    resolved_mode, layout = _resolve_docx_mode(mode, input_path)
    logger.info(
        "pdf-to-docx mode=%s (requested=%s) pages=%d images=%d vectors=%d rotated=%s tables=%s cols=%d",
        resolved_mode,
        mode,
        layout.page_count,
        layout.image_count,
        layout.vector_line_count,
        layout.has_rotated_text,
        layout.has_complex_tables,
        layout.column_count,
    )

    if resolved_mode == "preserve_layout":
        _build_preserve_layout_docx(input_path, output_path, dpi=dpi)
        return ENGINE_PRESERVE, resolved_mode, layout

    if resolved_mode == "fixed_layout":
        _build_fixed_layout_docx(input_path, output_path)
        return ENGINE_FIXED, resolved_mode, layout

    _build_editable_docx(input_path, output_path)
    return ENGINE_EDITABLE, resolved_mode, layout


# ---------------------------------------------------------------------------
# DOCX text extraction & translation replacement
# ---------------------------------------------------------------------------


def extract_docx_texts(docx_path: Path) -> list[dict[str, Any]]:
    """Extract text entries from DOCX paragraphs and table cells for translation."""
    doc = Document(str(docx_path))
    entries: list[dict[str, Any]] = []
    idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            entries.append({"id": idx, "text": text})
            idx += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    entries.append({"id": idx, "text": text})
                    idx += 1

    return entries


def apply_docx_translations(
    docx_path: Path,
    translations: dict[int, str],
    output_path: Path,
) -> None:
    """Replace text in DOCX with translations, preserving formatting."""
    doc = Document(str(docx_path))
    idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        translated = translations.get(idx)
        if translated:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = translated
            else:
                para.add_run(translated)
        idx += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue
                translated = translations.get(idx)
                if translated:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.text = ""
                    first_p = cell.paragraphs[0]
                    if first_p.runs:
                        first_p.runs[0].text = translated
                    else:
                        first_p.add_run(translated)
                idx += 1

    doc.save(str(output_path))
